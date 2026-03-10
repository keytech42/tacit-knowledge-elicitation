"""Tests for file parser service and upload endpoint."""

import io
import json
from unittest.mock import AsyncMock, patch

import pytest
from httpx import AsyncClient

from app.config import settings
from app.models.user import User
from app.services.file_parser import (
    DocxParser,
    JsonParser,
    PdfParser,
    TextParser,
    get_parser,
    parse_file,
)
from tests.conftest import auth_header

pytestmark = pytest.mark.asyncio


# ---------------------------------------------------------------------------
# Parser unit tests
# ---------------------------------------------------------------------------


class TestTextParser:
    def test_plain_text(self):
        parser = TextParser()
        result = parser.parse(b"Hello, world!", "test.txt")
        assert result == "Hello, world!"

    def test_utf8_text(self):
        parser = TextParser()
        content = "한국어 텍스트 테스트".encode("utf-8")
        result = parser.parse(content, "test.txt")
        assert "한국어" in result

    def test_markdown(self):
        parser = TextParser()
        md = b"# Heading\n\nParagraph text.\n\n- Item 1\n- Item 2"
        result = parser.parse(md, "test.md")
        assert "# Heading" in result
        assert "Item 1" in result


class TestPdfParser:
    def test_parse_pdf(self):
        """Create a minimal PDF in memory and parse it."""
        import pymupdf

        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Test PDF content for extraction.")
        pdf_bytes = doc.tobytes()
        doc.close()

        parser = PdfParser()
        result = parser.parse(pdf_bytes, "test.pdf")
        assert "Test PDF content" in result

    def test_multi_page_pdf(self):
        import pymupdf

        doc = pymupdf.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page {i + 1} content.")
        pdf_bytes = doc.tobytes()
        doc.close()

        parser = PdfParser()
        result = parser.parse(pdf_bytes, "test.pdf")
        assert "Page 1" in result
        assert "Page 2" in result
        assert "Page 3" in result


class TestDocxParser:
    def test_parse_docx(self):
        """Create a minimal DOCX in memory and parse it."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("First paragraph of the document.")
        doc.add_paragraph("Second paragraph with more content.")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        parser = DocxParser()
        result = parser.parse(docx_bytes, "test.docx")
        assert "First paragraph" in result
        assert "Second paragraph" in result

    def test_skips_empty_paragraphs(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Content paragraph.")
        doc.add_paragraph("")  # empty
        doc.add_paragraph("   ")  # whitespace only
        doc.add_paragraph("Another paragraph.")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        parser = DocxParser()
        result = parser.parse(docx_bytes, "test.docx")
        assert "Content paragraph." in result
        assert "Another paragraph." in result
        # Empty paragraphs should not produce blank lines
        assert "\n\n\n" not in result


class TestJsonParser:
    def test_simple_string(self):
        parser = JsonParser()
        data = json.dumps("just a string").encode()
        result = parser.parse(data, "test.json")
        assert result == "just a string"

    def test_flat_object(self):
        parser = JsonParser()
        data = json.dumps({"title": "Hello", "body": "World"}).encode()
        result = parser.parse(data, "test.json")
        assert "Hello" in result
        assert "World" in result

    def test_nested_object(self):
        parser = JsonParser()
        data = json.dumps({
            "doc": {"sections": [{"text": "Section 1"}, {"text": "Section 2"}]},
        }).encode()
        result = parser.parse(data, "test.json")
        assert "Section 1" in result
        assert "Section 2" in result

    def test_array_of_strings(self):
        parser = JsonParser()
        data = json.dumps(["one", "two", "three"]).encode()
        result = parser.parse(data, "test.json")
        assert "one" in result
        assert "two" in result
        assert "three" in result

    def test_ignores_non_string_values(self):
        parser = JsonParser()
        data = json.dumps({"count": 42, "flag": True, "text": "keep"}).encode()
        result = parser.parse(data, "test.json")
        assert "keep" in result
        assert "42" not in result

    def test_max_depth_protection(self):
        parser = JsonParser()
        # Build deeply nested structure
        data: dict = {"text": "found"}
        for _ in range(15):
            data = {"nested": data}
        content = json.dumps(data).encode()
        result = parser.parse(content, "test.json")
        # Should not crash, "found" may or may not be extracted depending on depth
        assert isinstance(result, str)


# ---------------------------------------------------------------------------
# parse_file / get_parser tests
# ---------------------------------------------------------------------------


class TestParseFile:
    def test_detect_by_content_type(self):
        parser = get_parser("text/plain", "unknown")
        assert isinstance(parser, TextParser)

    def test_detect_by_extension(self):
        parser = get_parser(None, "document.pdf")
        assert isinstance(parser, PdfParser)

    def test_extension_fallback_when_content_type_unknown(self):
        parser = get_parser("application/octet-stream", "file.json")
        # application/octet-stream not in registry, falls back to .json
        assert isinstance(parser, JsonParser)

    def test_unsupported_format_raises(self):
        with pytest.raises(ValueError, match="Unsupported"):
            get_parser(None, "data.xlsx")

    def test_file_too_large(self):
        content = b"x" * (10 * 1024 * 1024 + 1)  # 10MB + 1 byte
        with pytest.raises(ValueError, match="too large"):
            parse_file(content, "big.txt")

    def test_parse_text_file(self):
        result = parse_file(b"Hello text", "doc.txt")
        assert result == "Hello text"

    def test_parse_markdown_by_extension(self):
        result = parse_file(b"# Title\n\nBody", "readme.md")
        assert "# Title" in result


# ---------------------------------------------------------------------------
# Upload endpoint integration tests
# ---------------------------------------------------------------------------


class TestExtractFromFileEndpoint:
    async def test_upload_text_file(self, client: AsyncClient, admin_user: User, db):
        mock_response = {"task_id": "upload-task-1", "status": "accepted"}
        with patch("app.services.worker_client.trigger_extract_questions", new_callable=AsyncMock, return_value=mock_response), \
             patch.object(settings, "WORKER_URL", "http://worker:8001"):
            r = await client.post(
                "/api/v1/ai/extract-from-file",
                files={"file": ("test.txt", b"Document content for extraction.", "text/plain")},
                data={"document_title": "My Doc", "domain": "testing", "max_questions": "5"},
                headers=auth_header(admin_user),
            )

        assert r.status_code == 200
        data = r.json()
        assert data["task_type"] == "extract_questions"
        assert data["worker_task_id"] == "upload-task-1"
        assert data["status"] == "running"

    async def test_upload_uses_filename_as_title(self, client: AsyncClient, admin_user: User, db):
        mock_response = {"task_id": "upload-task-2", "status": "accepted"}
        with patch("app.services.worker_client.trigger_extract_questions", new_callable=AsyncMock, return_value=mock_response) as mock_trigger, \
             patch.object(settings, "WORKER_URL", "http://worker:8001"):
            r = await client.post(
                "/api/v1/ai/extract-from-file",
                files={"file": ("report.txt", b"Content here.", "text/plain")},
                data={},  # no document_title
                headers=auth_header(admin_user),
            )

        assert r.status_code == 200
        # The trigger should use filename as title
        call_kwargs = mock_trigger.call_args
        assert call_kwargs.kwargs.get("document_title") == "report.txt" or \
               (call_kwargs.args and "report.txt" in str(call_kwargs))

    async def test_upload_unsupported_format(self, client: AsyncClient, admin_user: User):
        with patch.object(settings, "WORKER_URL", "http://worker:8001"):
            r = await client.post(
                "/api/v1/ai/extract-from-file",
                files={"file": ("data.xlsx", b"fake xlsx content", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
                headers=auth_header(admin_user),
            )
        assert r.status_code == 400
        assert "Unsupported" in r.json()["detail"]

    async def test_upload_empty_file(self, client: AsyncClient, admin_user: User):
        with patch.object(settings, "WORKER_URL", "http://worker:8001"):
            r = await client.post(
                "/api/v1/ai/extract-from-file",
                files={"file": ("empty.txt", b"", "text/plain")},
                headers=auth_header(admin_user),
            )
        assert r.status_code == 400
        assert "no extractable text" in r.json()["detail"]

    async def test_upload_requires_admin(self, client: AsyncClient, author_user: User):
        r = await client.post(
            "/api/v1/ai/extract-from-file",
            files={"file": ("test.txt", b"Content", "text/plain")},
            headers=auth_header(author_user),
        )
        assert r.status_code == 403

    async def test_upload_requires_worker(self, client: AsyncClient, admin_user: User):
        with patch.object(settings, "WORKER_URL", ""):
            r = await client.post(
                "/api/v1/ai/extract-from-file",
                files={"file": ("test.txt", b"Content", "text/plain")},
                headers=auth_header(admin_user),
            )
        assert r.status_code == 503
